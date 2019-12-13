# -*- coding: utf-8 -*-
"""
Created on Mon Oct 21 19:28:34 2019

@author: Imsourav
"""

import shutil
from PIL import Image
import glob
from docx.shared import Inches
from docx import Document
import os
from flowchart_generator import *
def dell12(num1,output):    
    document = Document()
    
    document.add_heading('Hand Drawn', 0)
    
    #document.add_picture('flowchart/output/output.jpg', width=Inches(1.25))
    cnt=0
    for filename in glob.glob('flowchart/store/*.jpg'): 
        cnt=cnt+1
    print(cnt) 
    print(num1)
    temp=2
    temp= cnt - num1    
    for filename in glob.glob('flowchart/store/*.jpg'): 
        temp=temp-1
        if(temp>=0):
            continue
        im = Image.open(filename)
        
        imgwidth, imgheight = im.size
        #print(imgwidth)
        #print(imgheight)
        p = document.add_paragraph()
        runner = p.add_run()
        if imgwidth<100:
            runner.add_text('\t\t')
        elif imgwidth<200:
            runner.add_text('\t')
        imgwidth/=200
        imgheight/=200
        
        runner.add_picture(filename,width=Inches(imgwidth),height=Inches(imgheight))
        text=myDict[filename]
        document.add_paragraph(text)
        #os.remove(filename)
    #output=output.replace('/', '\')    
    document.save(output)
    
    """
    st =  output
    st = st[:-4]
    st=st+"jpg"
    print(st)
    print(output)
    shutil.copy("output.jpg",st)
    """
#if __name__== "__main__" :
    #dell12(1)