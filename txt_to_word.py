
import json
import time
import sys
from docx.shared import Inches
from docx import Document
def split1(ok):
    cnt=0 
    fo = open('words_dictionary.json','r')
    data = json.load(fo)
    with open(ok,'r') as f:
        for line in f:
            word = line.split()
            for temp in word:
                searchKey=temp.lower()
                if temp.isdigit():
                    cnt=cnt+1
                elif searchKey in data.keys():
                    cnt=cnt+1
                    
    return cnt         
def write1(temp,output,val):
    if val==1:
        document = Document()
        document.add_heading('Hand Drawn', 0)
        
        #document.add_picture('flowchart/output/output.jpg', width=Inches(1.25)) 
        with open(temp,'r') as f:
            for line in f:
                   print(line)       
                   document.add_paragraph(line)
                   #document.add_paragraph('first item in ordered list', style='ListNumber')
        document.save(output)
    else:
        f1= open(output,"a")
        with open(temp,'r') as f:
            for line in f:
                print(line)                         
                f1.write(line)
                f1.write("\n")
        f1.close()
     
    
def dell(output,val):    
   #cnt1=split1('output/output1.txt')
   #cnt2=split1('output/output2.txt')
   cnt3=split1('output/output3.txt')
   cnt4=split1('output/output4.txt')
   print(cnt3,cnt4)
   if cnt4>=3 or cnt4>=cnt3:
       write1('output/output4.txt',output,val)
   else:
       write1('output/output3.txt',output,val)
   """
    elif cnt3>=cnt2 and cnt3>=cnt1 and cnt3>=cnt4:
       write1('output/output3.txt',output,val)
   else:
       write1('output/output4.txt',output,val)
    """